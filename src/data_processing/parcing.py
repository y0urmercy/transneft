import logging
from typing import List, Dict, Tuple
from datetime import datetime
from pathlib import Path
import re
import json
import pandas as pd
from docx import Document
import nltk
from nltk.tokenize import sent_tokenize
import torch
from transformers import pipeline, AutoTokenizer, AutoModel
import evaluate
from sentence_transformers import SentenceTransformer
import numpy as np
from sklearn.metrics.pairwise import cosine_similarity
import hashlib

# NLTK setup
nltk.download('punkt', quiet=True)

# Logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

class AdvancedTransneftParser:
    def __init__(self, docx_path: str):
        self.docx_path = docx_path
        self.document = None
        self.sections = {}
        self.qa_triplets = []
        self.benchmark_data = []
        
        # Initialize models for quality assessment
        try:
            self.embedder = SentenceTransformer('intfloat/multilingual-e5-large')
            self.rouge = evaluate.load('rouge')
            logger.info("Quality assessment models loaded successfully")
        except Exception as e:
            logger.warning(f"Could not load quality models: {e}")
            self.embedder = None
            self.rouge = None

    def load_document(self) -> bool:
        """Load and validate DOCX document"""
        try:
            self.document = Document(self.docx_path)
            if len(self.document.paragraphs) == 0:
                logger.error("Document appears to be empty")
                return False
            logger.info(f"Document loaded successfully: {len(self.document.paragraphs)} paragraphs")
            return True
        except Exception as e:
            logger.error(f"Failed to load document: {e}")
            return False

    def extract_document_structure(self) -> Dict:
        """Advanced document structure extraction with hierarchical parsing"""
        sections = {}
        current_section = "Введение"
        current_content = []
        hierarchy_level = 0
        last_heading_level = 0
        
        for i, paragraph in enumerate(self.document.paragraphs):
            text = paragraph.text.strip()
            if not text:
                continue
                
            # Detect heading with improved logic
            heading_level = self._detect_heading_level(paragraph, text)
            
            if heading_level > 0:
                # Save previous section
                if current_content and self._is_valid_content(current_content):
                    content_text = ' '.join(current_content)
                    section_id = self._generate_section_id(current_section)
                    
                    sections[section_id] = {
                        'title': current_section,
                        'content': content_text,
                        'paragraphs': len(current_content),
                        'word_count': len(content_text.split()),
                        'heading_level': last_heading_level,
                        'entities': self.extract_advanced_entities(content_text),
                        'content_hash': hashlib.md5(content_text.encode()).hexdigest()
                    }
                
                current_section = text
                current_content = []
                last_heading_level = heading_level
            else:
                if len(text) > 10:  # Filter out very short paragraphs
                    current_content.append(text)
        
        # Save final section
        if current_content and self._is_valid_content(current_content):
            content_text = ' '.join(current_content)
            section_id = self._generate_section_id(current_section)
            
            sections[section_id] = {
                'title': current_section,
                'content': content_text,
                'paragraphs': len(current_content),
                'word_count': len(content_text.split()),
                'heading_level': last_heading_level,
                'entities': self.extract_advanced_entities(content_text),
                'content_hash': hashlib.md5(content_text.encode()).hexdigest()
            }
        
        self.sections = sections
        logger.info(f"Extracted {len(sections)} structured sections")
        return sections

    def _detect_heading_level(self, paragraph, text: str) -> int:
        """Detect heading level with multiple strategies"""
        # Strategy 1: Check paragraph style
        if hasattr(paragraph, 'style') and paragraph.style.name:
            style_name = paragraph.style.name.lower()
            if 'heading' in style_name:
                try:
                    return int(paragraph.style.name.split()[-1])
                except:
                    return 1
        
        # Strategy 2: Text-based detection
        text_lower = text.lower()
        
        # Level 1: Main headings (usually short, no punctuation)
        if (len(text) < 100 and 
            not text.endswith(('.', '!', '?', ':', ';')) and
            (text.isupper() or 
             any(word in text_lower for word in ['раздел', 'глава', 'часть', 'проект', 'история', 'факты']) or
             re.match(r'^\d+[\.\)]', text) or
             re.match(r'^[IVXLCDM]+\.', text))):
            return 1
        
        # Level 2: Subheadings
        if (len(text) < 200 and
            not text.endswith(('.', '!', '?')) and
            (text[0].isupper() and not text.isupper())):
            return 2
        
        return 0

    def _is_valid_content(self, content: List[str]) -> bool:
        """Validate if content is meaningful"""
        full_text = ' '.join(content)
        return len(full_text) >= 50 and len(full_text.split()) >= 10

    def _generate_section_id(self, title: str) -> str:
        """Generate unique section ID"""
        clean_title = re.sub(r'[^\w\s-]', '', title.lower())
        clean_title = re.sub(r'[-\s]+', '_', clean_title)
        return clean_title[:50]

    def extract_advanced_entities(self, text: str) -> Dict[str, List[str]]:
        """Extract comprehensive entities for Transneft context"""
        entities = {
            'dates': [],
            'years': [],
            'financials': [],
            'projects': [],
            'locations': [],
            'pipelines': [],
            'companies': [],
            'persons': []
        }
        
        # Dates
        dates = re.findall(r'\b\d{1,2}\.\d{1,2}\.\d{4}\b', text)
        entities['dates'].extend(dates)
        
        # Years
        years = re.findall(r'\b(?:19|20)\d{2}\b', text)
        entities['years'].extend(years)
        
        # Financial figures
        financials = re.findall(
            r'\b\d+(?:\s*\d{3})*(?:\,\d+)?\s*(?:млн|млрд|тыс|тонн|км|рублей|руб|акций|тонн\/год)(?:\s*\/\s*год)?\b', 
            text, 
            re.IGNORECASE
        )
        entities['financials'].extend(financials)
        
        # Projects
        projects = re.findall(
            r'(?:проект|нефтепровод|трубопровод|система|магистраль)\s+[«"]([^»"]+)[»"]',
            text
        )
        entities['projects'].extend(projects)
        
        # Pipelines and systems
        pipelines = re.findall(
            r'\b(?:ВСТО|БТС|КТК|Дружба|Заполярье|Куюмба|Тайшет|Сковородино|Козьмино|Приморск|Усть-Луга|Пурпе|Самотлор)\b',
            text
        )
        entities['pipelines'].extend(pipelines)
        
        # Locations
        locations = re.findall(
            r'\b(?:Москва|Санкт-Петербург|Новороссийск|Приморск|Усть-Луга|Тайшет|Красноярск|Иркутск|Хабаровск|Владивосток)\b',
            text
        )
        entities['locations'].extend(locations)
        
        # Companies
        companies = re.findall(
            r'\b(?:Транснефть|Роснефть|Лукойл|Газпром|Нобель|Совет\s+директоров|Правление)\b',
            text, re.IGNORECASE
        )
        entities['companies'].extend(companies)
        
        # Historical figures
        persons = re.findall(
            r'\b(?:Менделеев|Шухов|Щукин|Нобель|Путин)\b',
            text
        )
        entities['persons'].extend(persons)
        
        return entities

    def generate_qa_triplets_advanced(self) -> List[Dict]:
        """Generate high-quality QA triplets using multiple strategies"""
        self.qa_triplets = []
        
        for section_id, section_data in self.sections.items():
            content = section_data['content']
            title = section_data['title']
            entities = section_data['entities']
            
            if len(content) < 200:  # Skip very short sections
                continue
            
            # Generate diverse question types
            question_sets = [
                self._generate_factual_questions(title, content, entities),
                self._generate_temporal_questions(title, content, entities),
                self._generate_comparative_questions(title, content, entities),
                self._generate_causal_questions(title, content, entities),
                self._generate_descriptive_questions(title, content, entities)
            ]
            
            all_questions = []
            for question_set in question_sets:
                all_questions.extend(question_set)
            
            # Create QA pairs with quality validation
            for i, (question, answer) in enumerate(all_questions[:8]):  # Limit per section
                if self._validate_qa_pair(question, answer, content):
                    triplet = {
                        'id': f"{section_id}_q{i}",
                        'section_id': section_id,
                        'section_title': title,
                        'context': content,
                        'question': question,
                        'answer': answer,
                        'context_length': len(content),
                        'word_count': len(content.split()),
                        'entities': entities,
                        'quality_score': self._calculate_qa_quality(question, answer, content),
                        'generation_timestamp': datetime.now().isoformat()
                    }
                    self.qa_triplets.append(triplet)
        
        logger.info(f"Generated {len(self.qa_triplets)} high-quality QA triplets")
        return self.qa_triplets

    def _generate_factual_questions(self, title: str, content: str, entities: Dict) -> List[Tuple[str, str]]:
        """Generate factual questions"""
        questions = []
        
        # Date-based questions
        for date in entities['dates'][:2]:
            question = f"Какие события связаны с датой {date} в разделе '{title}'?"
            answer = self._extract_answer_for_date(content, date)
            if answer:
                questions.append((question, answer))
        
        # Financial questions
        for financial in entities['financials'][:2]:
            question = f"Что означает показатель '{financial}' в контексте раздела '{title}'?"
            answer = self._extract_financial_context(content, financial)
            if answer:
                questions.append((question, answer))
        
        # Project questions
        for project in entities['projects'][:2]:
            question = f"В чем заключается проект '{project}', упомянутый в разделе '{title}'?"
            answer = self._extract_project_description(content, project)
            if answer:
                questions.append((question, answer))
        
        return questions

    def _generate_temporal_questions(self, title: str, content: str, entities: Dict) -> List[Tuple[str, str]]:
        """Generate temporal questions"""
        questions = []
        
        years = entities['years']
        if len(years) >= 2:
            question = f"Какие ключевые события произошли в период с {min(years)} по {max(years)} год согласно разделу '{title}'?"
            answer = self._extract_temporal_period(content, min(years), max(years))
            if answer:
                questions.append((question, answer))
        
        # Historical development questions
        if any(year in content for year in ['1993', '2000', '2010', '2020']):
            question = f"Как развивалась компания Транснефть по данным раздела '{title}'?"
            answer = self._extract_development_timeline(content)
            if answer:
                questions.append((question, answer))
        
        return questions

    def _generate_comparative_questions(self, title: str, content: str, entities: Dict) -> List[Tuple[str, str]]:
        """Generate comparative questions"""
        questions = []
        
        pipelines = entities['pipelines']
        if len(pipelines) >= 2:
            question = f"Какие сравнения можно провести между проектами {', '.join(pipelines[:2])} согласно разделу '{title}'?"
            answer = self._extract_comparison(content, pipelines[:2])
            if answer:
                questions.append((question, answer))
        
        return questions

    def _generate_causal_questions(self, title: str, content: str, entities: Dict) -> List[Tuple[str, str]]:
        """Generate causal questions"""
        questions = []
        
        if any(word in content.lower() for word in ['причина', 'цель', 'задача', 'основание']):
            question = f"Какие цели и задачи преследуются в рамках раздела '{title}'?"
            answer = self._extract_goals_and_objectives(content)
            if answer:
                questions.append((question, answer))
        
        return questions

    def _generate_descriptive_questions(self, title: str, content: str, entities: Dict) -> List[Tuple[str, str]]:
        """Generate descriptive questions"""
        questions = []
        
        # General descriptive questions
        base_questions = [
            (f"О чем говорится в разделе '{title}'?", self._extract_summary(content)),
            (f"Какая основная информация представлена в разделе '{title}'?", self._extract_main_points(content)),
            (f"Какие ключевые аспекты освещаются в разделе '{title}'?", self._extract_key_aspects(content))
        ]
        
        questions.extend([q for q in base_questions if q[1]])
        
        return questions

    def _extract_answer_for_date(self, content: str, date: str) -> str:
        """Extract context around specific date"""
        sentences = sent_tokenize(content)
        for sentence in sentences:
            if date in sentence:
                return sentence
        return self._extract_relevant_sentence(content, date)

    def _extract_financial_context(self, content: str, financial: str) -> str:
        """Extract context for financial figures"""
        sentences = sent_tokenize(content)
        for sentence in sentences:
            if financial in sentence:
                return sentence
        return ""

    def _extract_project_description(self, content: str, project: str) -> str:
        """Extract project description"""
        # Look for sentences that describe the project
        sentences = sent_tokenize(content)
        project_sentences = []
        
        for sentence in sentences:
            if project in sentence and len(sentence) > 20:
                project_sentences.append(sentence)
        
        return ' '.join(project_sentences[:2]) if project_sentences else ""

    def _extract_temporal_period(self, content: str, start_year: str, end_year: str) -> str:
        """Extract information about temporal period"""
        relevant_sentences = []
        sentences = sent_tokenize(content)
        
        for sentence in sentences:
            if any(year in sentence for year in [start_year, end_year]):
                relevant_sentences.append(sentence)
        
        return ' '.join(relevant_sentences[:3]) if relevant_sentences else ""

    def _extract_development_timeline(self, content: str) -> str:
        """Extract development timeline"""
        timeline_keywords = ['основан', 'создан', 'начало', 'развитие', 'расширение', 'строительство']
        relevant_sentences = []
        sentences = sent_tokenize(content)
        
        for sentence in sentences:
            if any(keyword in sentence.lower() for keyword in timeline_keywords):
                relevant_sentences.append(sentence)
        
        return ' '.join(relevant_sentences[:4]) if relevant_sentences else ""

    def _extract_comparison(self, content: str, items: List[str]) -> str:
        """Extract comparison between items"""
        comparison_indicators = ['сравнен', 'отлич', 'преимуществ', 'особенност']
        relevant_sentences = []
        sentences = sent_tokenize(content)
        
        for sentence in sentences:
            if (any(item in sentence for item in items) and 
                any(indicator in sentence.lower() for indicator in comparison_indicators)):
                relevant_sentences.append(sentence)
        
        return ' '.join(relevant_sentences[:2]) if relevant_sentences else ""

    def _extract_goals_and_objectives(self, content: str) -> str:
        """Extract goals and objectives"""
        goal_keywords = ['цель', 'задача', 'основание', 'причина']
        relevant_sentences = []
        sentences = sent_tokenize(content)
        
        for sentence in sentences:
            if any(keyword in sentence.lower() for keyword in goal_keywords):
                relevant_sentences.append(sentence)
        
        return ' '.join(relevant_sentences[:3]) if relevant_sentences else ""

    def _extract_summary(self, content: str) -> str:
        """Extract summary of content"""
        sentences = sent_tokenize(content)
        return ' '.join(sentences[:2]) if len(sentences) >= 2 else content[:300] + "..."

    def _extract_main_points(self, content: str) -> str:
        """Extract main points"""
        # Look for sentences with key information
        key_indicators = ['основн', 'ключев', 'главн', 'важн']
        relevant_sentences = []
        sentences = sent_tokenize(content)
        
        for sentence in sentences:
            if any(indicator in sentence.lower() for indicator in key_indicators):
                relevant_sentences.append(sentence)
        
        if relevant_sentences:
            return ' '.join(relevant_sentences[:3])
        else:
            return self._extract_summary(content)

    def _extract_key_aspects(self, content: str) -> str:
        """Extract key aspects"""
        return self._extract_main_points(content)

    def _extract_relevant_sentence(self, content: str, keyword: str) -> str:
        """Extract the most relevant sentence containing keyword"""
        sentences = sent_tokenize(content)
        for sentence in sentences:
            if keyword in sentence:
                return sentence
        return ""

    def _validate_qa_pair(self, question: str, answer: str, context: str) -> bool:
        """Validate QA pair quality"""
        if not question or not answer:
            return False
        
        if len(answer) < 10 or len(answer) > 1000:
            return False
        
        # Check if answer is relevant to context
        if answer not in context and not any(word in context for word in answer.split()[:3]):
            return False
        
        return True

    def _calculate_qa_quality(self, question: str, answer: str, context: str) -> float:
        """Calculate QA pair quality score"""
        score = 0.0
        
        # Length-based scoring
        if 50 <= len(answer) <= 500:
            score += 0.3
        
        # Question complexity
        if any(word in question.lower() for word in ['почему', 'как', 'сравнит', 'причин']):
            score += 0.2
        
        # Answer relevance (simple check)
        if any(word in context for word in answer.split()[:3]):
            score += 0.3
        
        # Entity coverage
        if len(question.split()) >= 5:
            score += 0.2
        
        return min(score, 1.0)

    def create_retriever_benchmark(self) -> List[Dict]:
        """Create benchmark for retriever evaluation"""
        self.benchmark_data = []
        
        for triplet in self.qa_triplets:
            benchmark_item = {
                'id': triplet['id'],
                'question': triplet['question'],
                'ground_truth_answer': triplet['answer'],
                'relevant_contexts': [{
                    'context': triplet['context'],
                    'section_title': triplet['section_title'],
                    'context_id': triplet['section_id']
                }],
                'metadata': {
                    'section': triplet['section_title'],
                    'context_length': triplet['context_length'],
                    'quality_score': triplet['quality_score'],
                    'entities': triplet['entities']
                }
            }
            self.benchmark_data.append(benchmark_item)
        
        logger.info(f"Created retriever benchmark with {len(self.benchmark_data)} items")
        return self.benchmark_data

    def calculate_metrics(self) -> Dict:
        """Calculate comprehensive metrics for the benchmark"""
        if not self.qa_triplets:
            return {}
        
        metrics = {
            'basic_stats': {
                'total_qa_pairs': len(self.qa_triplets),
                'total_sections': len(self.sections),
                'avg_context_length': np.mean([t['context_length'] for t in self.qa_triplets]),
                'avg_answer_length': np.mean([len(t['answer']) for t in self.qa_triplets]),
                'avg_quality_score': np.mean([t['quality_score'] for t in self.qa_triplets])
            },
            'question_types': self._analyze_question_types(),
            'coverage_metrics': self._calculate_coverage_metrics()
        }
        
        # Calculate semantic metrics if models are available
        if self.embedder and self.rouge:
            try:
                semantic_metrics = self._calculate_semantic_metrics()
                metrics.update(semantic_metrics)
            except Exception as e:
                logger.warning(f"Could not calculate semantic metrics: {e}")
        
        return metrics

    def _analyze_question_types(self) -> Dict:
        """Analyze distribution of question types"""
        question_types = {
            'factual': 0,
            'temporal': 0,
            'comparative': 0,
            'causal': 0,
            'descriptive': 0
        }
        
        for triplet in self.qa_triplets:
            question = triplet['question'].lower()
            if any(word in question for word in ['когда', 'дата', 'год', 'сколько']):
                question_types['factual'] += 1
            elif any(word in question for word in ['период', 'развитие', 'история']):
                question_types['temporal'] += 1
            elif any(word in question for word in ['сравнен', 'отлич', 'преимуществ']):
                question_types['comparative'] += 1
            elif any(word in question for word in ['почему', 'причина', 'цель', 'задача']):
                question_types['causal'] += 1
            else:
                question_types['descriptive'] += 1
        
        return question_types

    def _calculate_coverage_metrics(self) -> Dict:
        """Calculate coverage metrics"""
        total_entities = 0
        covered_entities = 0
        
        for section_id, section_data in self.sections.items():
            entities = section_data['entities']
            for entity_type, entity_list in entities.items():
                total_entities += len(entity_list)
                
                # Check if entities are covered in questions
                for entity in entity_list:
                    if any(entity in triplet['question'] for triplet in self.qa_triplets):
                        covered_entities += 1
        
        coverage = covered_entities / total_entities if total_entities > 0 else 0
        
        return {
            'entity_coverage': coverage,
            'section_coverage': len([s for s in self.sections.values() if s['word_count'] > 200]) / len(self.sections),
            'total_entities': total_entities,
            'covered_entities': covered_entities
        }

    def _calculate_semantic_metrics(self) -> Dict:
        """Calculate semantic similarity metrics"""
        answers = [t['answer'] for t in self.qa_triplets]
        contexts = [t['context'] for t in self.qa_triplets]
        
        # Calculate ROUGE scores
        rouge_scores = self.rouge.compute(
            predictions=answers,
            references=contexts,
            use_stemmer=True
        )
        
        # Calculate semantic similarity
        answer_embeddings = self.embedder.encode(answers)
        context_embeddings = self.embedder.encode(contexts)
        
        similarities = []
        for i in range(len(answers)):
            sim = cosine_similarity(
                [answer_embeddings[i]], 
                [context_embeddings[i]]
            )[0][0]
            similarities.append(sim)
        
        avg_similarity = np.mean(similarities)
        
        return {
            'rouge_scores': rouge_scores,
            'semantic_similarity': float(avg_similarity),
            'similarity_std': float(np.std(similarities))
        }

    def save_benchmark(self, output_dir: str = "transneft_benchmark") -> Dict:
        """Save comprehensive benchmark with all required formats"""
        Path(output_dir).mkdir(exist_ok=True)
        
        # Create main dataframe
        df = pd.DataFrame(self.qa_triplets)
        
        # 1. Save QA triplets in CSV format
        df.to_csv(f"{output_dir}/qa_triplets.csv", index=False, encoding='utf-8')
        
        # 2. Save benchmark in JSONL format for evaluation
        with open(f"{output_dir}/benchmark.jsonl", 'w', encoding='utf-8') as f:
            for item in self.benchmark_data:
                f.write(json.dumps(item, ensure_ascii=False) + '\n')
        
        # 3. Save detailed benchmark in JSON format
        detailed_benchmark = {
            "metadata": {
                "company": "ПАО Транснефть",
                "source_file": self.docx_path,
                "creation_date": datetime.now().isoformat(),
                "total_qa_pairs": len(self.qa_triplets),
                "total_sections": len(self.sections),
                "benchmark_version": "1.0"
            },
            "sections": self.sections,
            "qa_triplets": self.qa_triplets,
            "retriever_benchmark": self.benchmark_data,
            "metrics": self.calculate_metrics()
        }
        
        with open(f"{output_dir}/detailed_benchmark.json", 'w', encoding='utf-8') as f:
            json.dump(detailed_benchmark, f, ensure_ascii=False, indent=2)
        
        # 4. Save contexts for retriever training
        contexts_data = []
        for section_id, section_data in self.sections.items():
            contexts_data.append({
                "context_id": section_id,
                "title": section_data['title'],
                "content": section_data['content'],
                "entities": section_data['entities'],
                "metadata": {
                    "word_count": section_data['word_count'],
                    "heading_level": section_data['heading_level']
                }
            })
        
        with open(f"{output_dir}/contexts.json", 'w', encoding='utf-8') as f:
            json.dump(contexts_data, f, ensure_ascii=False, indent=2)
        
        # 5. Save statistics report
        stats = self._generate_statistics_report()
        with open(f"{output_dir}/statistics.json", 'w', encoding='utf-8') as f:
            json.dump(stats, f, ensure_ascii=False, indent=2)
        
        logger.info(f"Benchmark saved successfully in {output_dir}")
        return stats

    def _generate_statistics_report(self) -> Dict:
        """Generate comprehensive statistics report"""
        metrics = self.calculate_metrics()
        
        report = {
            "summary": {
                "total_qa_pairs": len(self.qa_triplets),
                "total_sections": len(self.sections),
                "total_words": sum(t['word_count'] for t in self.qa_triplets),
                "benchmark_quality": "HIGH" if metrics['basic_stats']['avg_quality_score'] > 0.7 else "MEDIUM"
            },
            "quality_metrics": metrics,
            "section_analysis": {
                "sections_with_qa": len(set(t['section_id'] for t in self.qa_triplets)),
                "avg_qa_per_section": len(self.qa_triplets) / len(self.sections),
                "largest_section": max(self.sections.items(), key=lambda x: x[1]['word_count'])[0],
                "most_productive_section": max(
                    [(section_id, len([t for t in self.qa_triplets if t['section_id'] == section_id])) 
                     for section_id in self.sections.keys()], 
                    key=lambda x: x[1]
                )[0]
            }
        }
        
        return report

    def run_complete_pipeline(self) -> Dict:
        """Run complete parsing and benchmark creation pipeline"""
        logger.info("Starting complete Transneft benchmark pipeline...")
        
        if not self.load_document():
            raise Exception("Failed to load document")
        
        self.extract_document_structure()
        self.generate_qa_triplets_advanced()
        self.create_retriever_benchmark()
        
        if not self.qa_triplets:
            logger.warning("No QA triplets were generated!")
            return {}
        
        stats = self.save_benchmark()
        
        logger.info("Pipeline completed successfully!")
        return stats

# Example usage
if __name__ == "__main__":
    parser = AdvancedTransneftParser("Реестр данных о компании ПАО Транснефть для хакатона весна-лето 2026.docx")
    
    try:
        results = parser.run_complete_pipeline()
        print("\n" + "="*80)
        print("ТРАНСНЕФТЬ БЕНЧМАРК - УСПЕШНО СОЗДАН")
        print("="*80)
        print(f"📊 Создано QA пар: {results['summary']['total_qa_pairs']}")
        print(f"📁 Обработано разделов: {results['summary']['total_sections']}")
        print(f"🎯 Качество бенчмарка: {results['summary']['benchmark_quality']}")
        print(f"📈 Покрытие сущностей: {results['quality_metrics']['coverage_metrics']['entity_coverage']:.1%}")
        print("\nФайлы сохранены в папке: transneft_benchmark/")
        
    except Exception as e:
        logger.error(f"Pipeline failed: {e}")
        raise
