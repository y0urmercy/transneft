# docx_parser_fixed.py
from docx import Document
import pandas as pd
import json
import re
from pathlib import Path
import logging
from typing import List, Dict
from datetime import datetime

# Настройка логирования
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)


class TransneftDocxParser:
    def __init__(self, docx_path: str):
        self.docx_path = docx_path
        self.document = None
        self.sections = {}
        self.qa_triplets = []

    def load_document(self):
        """Загрузка DOCX документа"""
        try:
            self.document = Document(self.docx_path)
            logger.info(f"Документ загружен: {self.docx_path}")
            logger.info(f"Количество параграфов: {len(self.document.paragraphs)}")
            return True
        except Exception as e:
            logger.error(f"Ошибка загрузки документа: {e}")
            return False

    def extract_structure(self):
        """Извлечение структуры документа"""
        sections = {}
        current_section = "Введение"
        current_content = []

        for i, paragraph in enumerate(self.document.paragraphs):
            text = paragraph.text.strip()
            if not text:
                continue

            # Определяем заголовки по стилю или содержанию
            if self._is_heading(paragraph, text):
                # Сохраняем предыдущий раздел
                if current_content and len(' '.join(current_content)) > 50:
                    content_text = ' '.join(current_content)
                    sections[current_section] = {
                        'content': content_text,
                        'paragraphs': len(current_content),
                        'word_count': len(content_text.split())
                    }

                current_section = text
                current_content = []
            else:
                current_content.append(text)

        # Сохраняем последний раздел
        if current_content and len(' '.join(current_content)) > 50:
            content_text = ' '.join(current_content)
            sections[current_section] = {
                'content': content_text,
                'paragraphs': len(current_content),
                'word_count': len(content_text.split())
            }

        self.sections = sections
        logger.info(f"Извлечено разделов: {len(sections)}")
        return sections

    def _is_heading(self, paragraph, text: str) -> bool:
        """Определяет, является ли параграф заголовком"""
        # По стилю
        if hasattr(paragraph, 'style') and paragraph.style.name.startswith('Heading'):
            return True

        # По содержанию (короткий текст, без пунктуации в конце)
        if len(text) < 200 and not text.endswith(('.', '!', '?')):
            # Проверяем, что это не просто короткое предложение
            if (text.isupper() or
                    any(word in text.lower() for word in ['раздел', 'глава', 'часть', '§']) or
                    re.match(r'^\d+[\.\)]', text) or  # "1.", "1)"
                    re.match(r'^[IVXLCDM]+\.', text) or  # Римские цифры
                    re.match(r'^[А-Я][а-я]+\s*$', text)):  # Одно слово с заглавной буквы
                return True

        return False

    def extract_entities(self, text: str) -> List[Dict]:
        """Извлечение ключевых сущностей для Транснефть"""
        entities = []

        # Даты
        dates = re.findall(r'\b\d{1,2}\.\d{1,2}\.\d{4}\b', text)
        for date in dates:
            entities.append({'type': 'date', 'value': date})

        # Годы
        years = re.findall(r'\b(?:19|20)\d{2}\b', text)
        for year in years:
            entities.append({'type': 'year', 'value': year})

        # Финансовые показатели
        financials = re.findall(
            r'\b\d+(?:\s*\d{3})*(?:\,\d+)?\s*(?:млн|млрд|тыс|тонн|км|рублей|руб|акций)(?:\s*\/\s*год)?\b',
            text,
            re.IGNORECASE
        )
        for financial in financials:
            entities.append({'type': 'financial', 'value': financial})

        # Проекты Транснефть
        projects = re.findall(
            r'(?:проект|нефтепровод|трубопровод|система)\s+[«"]([^»"]+)[»"]',
            text
        )
        for project in projects:
            entities.append({'type': 'project', 'value': project})

        # Ключевые объекты
        objects = re.findall(
            r'\b(?:ВСТО|БТС|КТК|Заполярье|Куюмба|Тайшет|Сковородино|Козьмино|Приморск|Усть-Луга)\b',
            text
        )
        for obj in objects:
            entities.append({'type': 'object', 'value': obj})

        return entities

    def generate_section_questions(self, section_title: str, content: str, entities: List[Dict]) -> List[str]:
        """Генерация вопросов для раздела"""
        questions = []

        content_lower = content.lower()
        section_lower = section_title.lower()

        # Вопросы на основе типа раздела
        if any(word in section_lower for word in ['уставный', 'капитал', 'акци']):
            questions.extend([
                "Каков размер уставного капитала ПАО Транснефть?",
                "Из каких акций состоит уставный капитал?",
                "Как изменялся уставный капитал компании?",
            ])

        elif any(word in section_lower for word in ['истори', 'основан']):
            questions.extend([
                "Когда была основана компания Транснефть?",
                "Какие ключевые исторические вехи в развитии компании?",
                "Как происходило становление компании?",
            ])

        elif any(word in section_lower for word in ['проект', 'трубопровод', 'нефтепровод']):
            project_entities = [e['value'] for e in entities if e['type'] == 'project']
            if project_entities:
                for project in project_entities[:2]:
                    questions.extend([
                        f"В чем заключается проект '{project}'?",
                        f"Какие цели у проекта '{project}'?",
                        f"Каковы сроки реализации проекта '{project}'?",
                    ])
            else:
                questions.extend([
                    "О каком проекте идет речь в этом разделе?",
                    "Какие цели и задачи у этого проекта?",
                    "Каковы ключевые характеристики проекта?",
                ])

        elif any(word in content_lower for word in ['директор', 'правление', 'совет']):
            questions.extend([
                "Как устроена система управления в ПАО Транснефть?",
                "Какие органы управления существуют в компании?",
                "Кто входит в состав Совета директоров?",
            ])

        # Общие вопросы
        questions.extend([
            f"Какая информация содержится в разделе '{section_title}'?",
            f"О чем говорится в разделе '{section_title}'?",
            f"Какие ключевые факты представлены в разделе '{section_title}'?",
        ])

        return list(set(questions))[:4]  # Убираем дубли и ограничиваем 4 вопросами

    def create_qa_triplets(self):
        """Создание QA триплетов"""
        self.qa_triplets = []

        for section_title, section_data in self.sections.items():
            content = section_data['content']

            if len(content) < 100:  # Пропускаем короткие разделы
                continue

            entities = self.extract_entities(content)
            questions = self.generate_section_questions(section_title, content, entities)

            for i, question in enumerate(questions):
                # Создаем более конкретный ответ на основе вопроса
                answer = self._generate_targeted_answer(content, question)

                self.qa_triplets.append({
                    'section': section_title,
                    'context': content,
                    'question': question,
                    'answer': answer,
                    'question_id': f"{section_title}_{i}",
                    'entities': [e['value'] for e in entities],
                    'context_length': len(content),
                    'word_count': len(content.split()),
                })

        logger.info(f"Создано QA триплетов: {len(self.qa_triplets)}")

    def _generate_targeted_answer(self, context: str, question: str) -> str:
        """Генерирует целевой ответ на основе вопроса"""
        question_lower = question.lower()

        # Для вопросов о датах
        if any(word in question_lower for word in ['когда', 'дата', 'год', 'срок']):
            dates = re.findall(r'\b\d{1,2}\.\d{1,2}\.\d{4}\b', context)
            if dates:
                return f"В контексте упоминаются даты: {', '.join(dates)}. {context[:300]}..."

        # Для финансовых вопросов
        if any(word in question_lower for word in ['капитал', 'акци', 'финанс', 'размер']):
            financials = re.findall(
                r'\b\d+(?:\s*\d{3})*\s*(?:акций|рублей|млн|млрд|тонн|км)\b',
                context
            )
            if financials:
                return f"Финансовые показатели: {', '.join(financials)}. {context[:300]}..."

        # Для проектных вопросов
        if any(word in question_lower for word in ['проект', 'трубопровод', 'нефтепровод']):
            projects = re.findall(
                r'(?:проект|нефтепровод|трубопровод)\s+[«"]([^»"]+)[»"]',
                context
            )
            if projects:
                return f"Упомянутые проекты: {', '.join(projects)}. {context[:300]}..."

        # Возвращаем осмысленный отрывок
        sentences = re.split(r'[.!?]+', context)
        if len(sentences) > 2:
            return sentences[0] + ". " + sentences[1] + "."
        else:
            return context[:400] + "..." if len(context) > 400 else context

    def parse(self) -> pd.DataFrame:
        """Основной метод парсинга"""
        logger.info("Начало парсинга DOCX документа ПАО Транснефть...")

        if not self.load_document():
            raise Exception("Не удалось загрузить документ")

        self.extract_structure()
        self.create_qa_triplets()

        if not self.qa_triplets:
            logger.warning("Не создано ни одного QA триплета!")
            return pd.DataFrame()

        df = pd.DataFrame(self.qa_triplets)

        # Добавляем мета-информацию
        df['source_file'] = self.docx_path
        df['parse_timestamp'] = datetime.now().isoformat()

        logger.info(f"Парсинг завершен. Создано {len(df)} QA пар")
        return df

    def convert_to_serializable(self, obj):
        """Рекурсивно преобразует объекты в сериализуемые"""
        if isinstance(obj, (int, float, str, bool, type(None))):
            return obj
        elif isinstance(obj, (pd.Timestamp, datetime)):
            return obj.isoformat()
        elif isinstance(obj, (pd.Series, pd.DataFrame)):
            return obj.to_dict()
        elif isinstance(obj, dict):
            return {key: self.convert_to_serializable(value) for key, value in obj.items()}
        elif isinstance(obj, (list, tuple)):
            return [self.convert_to_serializable(item) for item in obj]
        elif hasattr(obj, 'dtype'):  # numpy/pandas типы
            if 'int' in str(obj.dtype):
                return int(obj)
            elif 'float' in str(obj.dtype):
                return float(obj)
            else:
                return str(obj)
        else:
            return str(obj)

    def save_benchmark(self, output_dir: str = "full_benchmark"):
        """Сохранение бенчмарка в разных форматах"""
        Path(output_dir).mkdir(exist_ok=True)

        df = self.parse()

        if df.empty:
            logger.error("Нет данных для сохранения!")
            return

        # 1. CSV для анализа
        df.to_csv(f"{output_dir}/qa_pairs.csv", index=False, encoding='utf-8')

        # 2. JSON для RAG системы (с преобразованием типов)
        benchmark = {
            "metadata": {
                "company": "ПАО Транснефть",
                "source_file": self.docx_path,
                "total_qa_pairs": int(len(df)),
                "total_sections": int(len(self.sections)),
                "avg_context_length": float(df['context_length'].mean()),
                "total_words": int(df['word_count'].sum()),
                "parse_date": datetime.now().isoformat()
            },
            "sections": list(self.sections.keys()),
            "qa_pairs": self.convert_to_serializable(df.to_dict('records'))
        }

        # Дополнительная проверка на сериализуемость
        benchmark_serializable = self.convert_to_serializable(benchmark)

        with open(f"{output_dir}/transneft_benchmark.json", 'w', encoding='utf-8') as f:
            json.dump(benchmark_serializable, f, ensure_ascii=False, indent=2)

        # 3. Текстовый файл с контекстами
        with open(f"{output_dir}/contexts.txt", 'w', encoding='utf-8') as f:
            for _, row in df.iterrows():
                f.write(f"РАЗДЕЛ: {row['section']}\n")
                f.write(f"КОНТЕКСТ: {row['context']}\n")
                f.write(f"ВОПРОС: {row['question']}\n")
                f.write(f"ОТВЕТ: {row['answer']}\n")
                f.write("=" * 100 + "\n\n")

        # 4. Статистика
        stats = self._calculate_statistics(df)
        stats_serializable = self.convert_to_serializable(stats)

        with open(f"{output_dir}/statistics.json", 'w', encoding='utf-8') as f:
            json.dump(stats_serializable, f, ensure_ascii=False, indent=2)

        self._print_statistics(stats, df)

        return df

    def _calculate_statistics(self, df: pd.DataFrame) -> Dict:
        """Расчет статистики"""
        if df.empty:
            return {}

        return {
            "total_qa_pairs": int(len(df)),
            "total_sections": int(len(self.sections)),
            "avg_context_length": float(df['context_length'].mean()),
            "avg_word_count": float(df['word_count'].mean()),
            "total_words": int(df['word_count'].sum()),
            "section_distribution": {
                str(section): int(len([t for t in self.qa_triplets if t['section'] == section]))
                for section in self.sections.keys()
            }
        }

    def _print_statistics(self, stats: Dict, df: pd.DataFrame):
        """Вывод статистики"""
        print("\n" + "=" * 70)
        print("СТАТИСТИКА ПАРСИНГА DOCX ДОКУМЕНТА ПАО ТРАНСНЕФТЬ")
        print("=" * 70)
        print(f"📊 Всего разделов: {stats['total_sections']}")
        print(f"❓ Всего QA пар: {stats['total_qa_pairs']}")
        print(f"📝 Средняя длина контекста: {stats['avg_context_length']:.0f} символов")
        print(f"🔢 Всего слов: {stats['total_words']}")

        print(f"\n📂 Основные разделы документа:")
        for i, section in enumerate(list(self.sections.keys())[:8]):
            count = stats['section_distribution'].get(str(section), 0)
            print(f"   {i + 1}. {section} ({count} QA пар)")

        if self.qa_triplets:
            print(f"\n🧪 Примеры вопросов:")
            for i, triplet in enumerate(self.qa_triplets[:3]):
                print(f"   {i + 1}. {triplet['question']}")


# Альтернативная простая версия для гарантированной работы
def safe_docx_parser(docx_path: str, output_dir: str = "safe_benchmark"):
    """Безопасный парсер с минимальной обработкой ошибок"""
    import pandas as pd
    import json

    print(f"🛡️  Безопасный парсинг файла: {docx_path}")

    try:
        doc = Document(docx_path)
        qa_pairs = []

        current_section = "Общая информация"
        current_content = []

        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if not text:
                continue

            # Простое определение заголовков
            if (len(text) < 100 and
                    not text.endswith(('.', '!', '?')) and
                    (text.isupper() or para.style.name.startswith('Heading') if hasattr(para, 'style') else False)):

                # Сохраняем предыдущий раздел
                if current_content:
                    content_text = ' '.join(current_content)
                    if len(content_text) > 100:
                        # Создаем QA пары
                        questions = [
                            f"Какая информация содержится в разделе '{current_section}'?",
                            f"О чем говорится в разделе '{current_section}'?",
                        ]

                        for q in questions:
                            qa_pairs.append({
                                'section': current_section,
                                'context': content_text,
                                'question': q,
                                'answer': content_text[:500] + '...' if len(content_text) > 500 else content_text
                            })

                current_section = text
                current_content = []
            else:
                if len(text) > 10:
                    current_content.append(text)

        # Сохраняем последний раздел
        if current_content:
            content_text = ' '.join(current_content)
            if len(content_text) > 100:
                questions = [
                    f"Какая информация содержится в разделе '{current_section}'?",
                    f"О чем говорится в разделе '{current_section}'?",
                ]

                for q in questions:
                    qa_pairs.append({
                        'section': current_section,
                        'context': content_text,
                        'question': q,
                        'answer': content_text[:500] + '...' if len(content_text) > 500 else content_text
                    })

        # Сохраняем результат
        Path(output_dir).mkdir(exist_ok=True)
        df = pd.DataFrame(qa_pairs)

        # CSV
        df.to_csv(f"{output_dir}/qa_pairs.csv", index=False, encoding='utf-8')

        # JSON с явным преобразованием типов
        benchmark = {
            "metadata": {
                "company": "ПАО Транснефть",
                "source_file": docx_path,
                "total_qa_pairs": int(len(df)),
                "parse_date": datetime.now().isoformat()
            },
            "qa_pairs": []
        }

        for _, row in df.iterrows():
            benchmark["qa_pairs"].append({
                "section": str(row['section']),
                "context": str(row['context']),
                "question": str(row['question']),
                "answer": str(row['answer'])
            })

        with open(f"{output_dir}/transneft_benchmark.json", 'w', encoding='utf-8') as f:
            json.dump(benchmark, f, ensure_ascii=False, indent=2)

        print(f"✅ УСПЕХ! Создано {len(df)} QA пар!")
        print(f"📁 Файлы сохранены в папке: {output_dir}")

        return df

    except Exception as e:
        print(f"❌ Ошибка: {e}")
        import traceback
        traceback.print_exc()
        return None


# Использование
if __name__ == "__main__":
    docx_file = "Реестр данных о компании ПАО Транснефть для хакатона весна-лета 2026.docx"

    print("Выберите вариант парсинга:")
    print("1. Безопасный парсинг (гарантированная работа)")
    print("2. Полный парсинг (расширенные функции)")

    choice = input("Введите 1 или 2: ").strip()

    if choice == "1":
        # Безопасный вариант
        safe_docx_parser(docx_file, "safe_output")
    else:
        # Полный вариант
        parser = TransneftDocxParser(docx_file)
        df = parser.save_benchmark("full_benchmark")

        print("\n🎉 Парсинг завершен! Созданы файлы:")
        print("   - full_benchmark/qa_pairs.csv")
        print("   - full_benchmark/transneft_benchmark.json")
        print("   - full_benchmark/contexts.txt")
        print("   - full_benchmark/statistics.json")
